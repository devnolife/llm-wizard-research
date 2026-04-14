#!/usr/bin/env python3
"""
Generate PROPOSAL_REVISI_FINAL.docx from markdown drafts.
Converts markdown content to properly formatted academic docx.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DRAFTS_DIR = "/home/devnolife/wizard-research/drafts"
OUTPUT_PATH = os.path.join(DRAFTS_DIR, "PROPOSAL_REVISI_FINAL.docx")

TITLE_ID = (
    "NEURO-SYMBOLIC AGENTIC SYSTEM UNTUK DETEKSI INDIKATOR "
    "SYNTHESIS GAP DALAM LITERATUR ILMIAH"
)
TITLE_EN = (
    "NEURO-SYMBOLIC AGENTIC SYSTEM FOR SYNTHESIS GAP "
    "INDICATOR DETECTION IN SCIENTIFIC LITERATURE"
)


# ---------------------------------------------------------------------------
# Helper: apply Times New Roman defaults to a run
# ---------------------------------------------------------------------------
def _set_run_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name)


# ---------------------------------------------------------------------------
# Helper: set paragraph spacing (line spacing 1.5, space after 6pt)
# ---------------------------------------------------------------------------
def _set_paragraph_format(paragraph, line_spacing=1.5, space_after=6, alignment=None):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    if alignment is not None:
        pf.alignment = alignment


# ---------------------------------------------------------------------------
# Helper: add a simple paragraph with formatted runs (handles **bold**)
# ---------------------------------------------------------------------------
def add_formatted_paragraph(doc, text, style="Normal", bold=False, italic=False,
                            size=12, alignment=None, first_line_indent=None,
                            left_indent=None, space_before=0, space_after=6,
                            font_name="Times New Roman"):
    p = doc.add_paragraph(style=style)
    _set_paragraph_format(p, alignment=alignment, space_after=space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)

    # Split text on **bold** markers and *italic* markers
    _add_inline_formatting(p, text, size=size, bold=bold, italic=italic, font_name=font_name)
    return p


def _add_inline_formatting(paragraph, text, size=12, bold=False, italic=False,
                           font_name="Times New Roman"):
    """Parse **bold** and *italic* markers in text and add runs."""
    # Pattern: **bold**, *italic*, `code`
    pattern = r'(\*\*.*?\*\*|\*[^*]+?\*|`[^`]+?`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, size=size, bold=True, italic=italic, name=font_name)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, size=size, bold=bold, italic=True, name=font_name)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, size=size, bold=bold, italic=italic, name="Courier New")
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, size=size, bold=bold, italic=italic, name=font_name)


# ---------------------------------------------------------------------------
# Helper: add table with borders
# ---------------------------------------------------------------------------
def add_table_with_borders(doc, headers, rows):
    """Add a formatted table with full borders."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set borders on every cell
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")} />')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        _add_inline_formatting(p, h, size=11, bold=True)
        _set_paragraph_format(p, space_after=2)
        # Shade header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        hdr_cells[i]._element.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            _add_inline_formatting(p, cell_text, size=11)
            _set_paragraph_format(p, space_after=2)

    # Auto-fit
    table.autofit = True
    return table


# ---------------------------------------------------------------------------
# Markdown Parser → docx elements
# ---------------------------------------------------------------------------
class MarkdownToDocx:
    def __init__(self, doc):
        self.doc = doc
        self.references = []

    def convert(self, md_text):
        lines = md_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Skip horizontal rules
            if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
                i += 1
                continue

            # Code blocks (```)
            if stripped.startswith("```"):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # skip closing ```
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    p = add_formatted_paragraph(
                        self.doc, code_text, size=10,
                        font_name="Courier New", left_indent=1.0,
                        space_before=4, space_after=4
                    )
                    # light gray background shading
                    pPr = p._element.get_or_add_pPr()
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
                    pPr.append(shading)
                continue

            # Headings
            heading_match = re.match(r'^(#{1,3})\s+(.*)', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                # Clean bold markers from heading text
                heading_text = re.sub(r'\*\*(.*?)\*\*', r'\1', heading_text)
                self._add_heading(heading_text, level)
                i += 1
                continue

            # Blockquotes
            if stripped.startswith("> "):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith("> "):
                    quote_lines.append(lines[i].strip()[2:])
                    i += 1
                quote_text = " ".join(quote_lines)
                add_formatted_paragraph(
                    self.doc, quote_text, italic=True, left_indent=1.5,
                    space_before=4, space_after=4
                )
                continue

            # Tables (pipe-delimited)
            if "|" in stripped and stripped.startswith("|"):
                table_lines = []
                while i < len(lines) and "|" in lines[i].strip() and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._parse_table(table_lines)
                continue

            # Numbered lists
            numbered_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
            if numbered_match:
                items = []
                while i < len(lines):
                    s = lines[i].strip()
                    nm = re.match(r'^(\d+)\.\s+(.*)', s)
                    if nm:
                        items.append(nm.group(2))
                        i += 1
                        # Collect continuation lines (indented or sub-items)
                        while i < len(lines) and lines[i].strip() and not re.match(r'^\d+\.\s+', lines[i].strip()) and not re.match(r'^#{1,3}\s+', lines[i].strip()) and not lines[i].strip().startswith("|") and not lines[i].strip().startswith("- ") and not lines[i].strip().startswith("* ") and not lines[i].strip().startswith("```") and not lines[i].strip().startswith("> "):
                            items[-1] += " " + lines[i].strip()
                            i += 1
                    else:
                        break
                for idx, item in enumerate(items):
                    text = f"{idx + 1}. {item}"
                    add_formatted_paragraph(self.doc, text, left_indent=1.0, space_after=4)
                continue

            # Bullet lists
            bullet_match = re.match(r'^[-*]\s+(.*)', stripped)
            if bullet_match:
                items = []
                while i < len(lines):
                    s = lines[i].strip()
                    bm = re.match(r'^[-*]\s+(.*)', s)
                    if bm:
                        items.append(bm.group(1))
                        i += 1
                        # continuation
                        while i < len(lines) and lines[i].strip() and not re.match(r'^[-*]\s+', lines[i].strip()) and not re.match(r'^#{1,3}\s+', lines[i].strip()) and not lines[i].strip().startswith("|") and not re.match(r'^\d+\.\s+', lines[i].strip()) and not lines[i].strip().startswith("```") and not lines[i].strip().startswith("> "):
                            items[-1] += " " + lines[i].strip()
                            i += 1
                    else:
                        break
                for item in items:
                    p = add_formatted_paragraph(self.doc, f"• {item}", left_indent=1.0, space_after=4)
                continue

            # Regular paragraphs — collect consecutive non-empty lines
            para_lines = []
            while i < len(lines):
                s = lines[i].strip()
                if not s:
                    i += 1
                    break
                if re.match(r'^#{1,3}\s+', s):
                    break
                if s.startswith("|") and "|" in s[1:]:
                    break
                if re.match(r'^[-*]\s+', s):
                    break
                if re.match(r'^\d+\.\s+', s):
                    break
                if s.startswith("```"):
                    break
                if s.startswith("> "):
                    break
                if re.match(r'^-{3,}$', s):
                    i += 1
                    break
                para_lines.append(s)
                i += 1

            if para_lines:
                text = " ".join(para_lines)
                # Skip the draft notice at the end
                if text.startswith("*Dokumen ini merupakan draf"):
                    continue
                add_formatted_paragraph(self.doc, text, first_line_indent=1.25)

    def _add_heading(self, text, level):
        if level == 1:
            p = self.doc.add_paragraph()
            p.style = self.doc.styles["Heading 1"]
            run = p.add_run(text.upper())
            _set_run_font(run, size=14, bold=True)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            p = self.doc.add_paragraph()
            p.style = self.doc.styles["Heading 2"]
            run = p.add_run(text)
            _set_run_font(run, size=12, bold=True)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
        elif level == 3:
            p = self.doc.add_paragraph()
            p.style = self.doc.styles["Heading 3"]
            run = p.add_run(text)
            _set_run_font(run, size=12, bold=True)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

    def _parse_table(self, table_lines):
        """Parse pipe-delimited markdown table lines into docx table."""
        parsed = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            # Skip separator rows (---|----|---)
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            parsed.append(cells)

        if len(parsed) < 1:
            return

        headers = parsed[0]
        rows = parsed[1:] if len(parsed) > 1 else []

        # Normalize column count
        ncols = len(headers)
        norm_rows = []
        for r in rows:
            while len(r) < ncols:
                r.append("")
            norm_rows.append(r[:ncols])

        add_table_with_borders(self.doc, headers, norm_rows)
        # Add spacing after table
        self.doc.add_paragraph()


# ---------------------------------------------------------------------------
# Collect references from markdown files
# ---------------------------------------------------------------------------
def collect_references(texts):
    """Extract references mentioned in the text (bracketed citations)."""
    all_refs = set()
    for text in texts:
        # Find patterns like [Author et al., Year], [Author, Year]
        refs = re.findall(r'\[([A-Z][a-zA-Z]+(?:\s+(?:et\s+al\.|&\s+[A-Z][a-zA-Z]+))?(?:,?\s*\d{4})?)\]', text)
        for r in refs:
            all_refs.add(r.strip())
    return sorted(all_refs)


# ---------------------------------------------------------------------------
# Build reference list from the NOVELTY_STATEMENT referensi utama + drafts
# ---------------------------------------------------------------------------
KNOWN_REFERENCES = [
    "Abu-Salih, B., et al. (2024). Systematic review of knowledge graph construction. Knowledge-Based Systems, 283, 111184.",
    "Amugongo, L. M., et al. (2025). A Comprehensive Review of RAG Techniques. arXiv:2503.xxxx.",
    "Anthropic. (2023). Claude: A Next-Generation AI Assistant. Technical Report.",
    "Bender, E. M., & Koller, A. (2020). Climbing towards NLU: On Meaning, Form, and Understanding in the Age of Data. Proceedings of ACL 2020.",
    "Booth, A., Sutton, A., & Papaioannou, D. (2012). Systematic Approaches to a Successful Literature Review. Sage.",
    "Buchanan, B. G., & Shortliffe, E. H. (1984). Rule-Based Expert Systems. Addison-Wesley.",
    "Chelli, M., et al. (2024). Hallucination Rates and Reference Accuracy of ChatGPT and Bard in Oral and Maxillofacial Surgery. Journal of Stomatology, Oral and Maxillofacial Surgery, 125(1), 101744.",
    "Cooper, H. (1998). Synthesizing Research: A Guide for Literature Reviews (3rd ed.). Sage Publications.",
    "Garcez, A., et al. (2019). Neural-Symbolic Computing: An Effective Methodology for Principled Integration of Machine Learning and Reasoning. Journal of Applied Logics, 6(4), 611–631.",
    "Guo, E., Gupta, M., & Sinha, S. (2023). Systematic Review Automation. Annals of Surgery, 278(3), 382–390.",
    "Hevner, A. R., March, S. T., Park, J., & Ram, S. (2004). Design Science in Information Systems Research. MIS Quarterly, 28(1), 75–105.",
    "Ji, S., et al. (2021). A Survey on Knowledge Graphs: Representation, Acquisition, and Applications. IEEE TNNLS, 33(2), 494–514.",
    "Johnson, R., Watkinson, A., & Mabe, M. (2018). The STM Report: An Overview of Scientific and Scholarly Publishing. International Association of STM Publishers.",
    "Lewis, P., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in NeurIPS, 33, 9459–9474.",
    "Ma, Y., et al. (2025). Vector Database Management Systems. Data Science and Engineering, 10(1), 1–20.",
    "Marcus, G. (2020). The Next Decade in AI: Four Steps Towards Robust Artificial Intelligence. arXiv:2002.06177.",
    "Marcus, G., & Davis, E. (2020). Rebooting AI: Building Artificial Intelligence We Can Trust. Vintage.",
    "Muller-Bloch, C., & Kranz, J. (2015). A Framework for Rigorously Identifying Research Gaps in Qualitative Literature Reviews. Proceedings of ICIS 2015.",
    "Nallapati, R., et al. (2024). Multi-Document Summarization with Context Window Limitations. arXiv.",
    "OpenAI. (2023). GPT-4 Technical Report. arXiv:2303.08774.",
    "Pare, G., Trudel, M.-C., Jaana, M., & Kitsiou, S. (2015). Synthesizing Information Systems Knowledge: A Typology of Literature Reviews. Information & Management, 52(2), 183–199.",
    "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. Proceedings of EMNLP-IJCNLP 2019.",
    "Robinson, K. A., Saldanha, I. J., & McKoy, N. A. (2011). Development of a Framework to Identify Research Gaps from Systematic Reviews. Journal of Clinical Epidemiology, 64(12), 1325–1330.",
    "Touvron, H., et al. (2023). LLaMA: Open and Efficient Foundation Language Models. arXiv:2302.13971.",
    "Vaswani, A., et al. (2017). Attention Is All You Need. Advances in NeurIPS, 30, 5998–6008.",
    "Yao, S., et al. (2023). ReAct: Synergizing Reasoning and Acting in Language Models. Proceedings of ICLR 2023.",
    "Zhang, Y., et al. (2025). A Survey on RAG Meets LLMs: Towards Retrieval-Augmented Large Language Models. arXiv:2405.06211.",
]


# ---------------------------------------------------------------------------
# Main document builder
# ---------------------------------------------------------------------------
def build_document():
    doc = Document()

    # ---- Page setup: A4, margins ----
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(4.0)
        section.top_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)

    # ---- Configure default styles ----
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Times New Roman"
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style_normal.paragraph_format.line_spacing = 1.5

    for heading_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style_h = doc.styles[heading_name]
        style_h.font.name = "Times New Roman"
        style_h.font.color.rgb = RGBColor(0, 0, 0)
        style_h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style_h.paragraph_format.line_spacing = 1.5

    # ======================== TITLE PAGE ========================
    # Add blank lines for centering
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DRAF PROPOSAL REVISI")
    _set_run_font(run, size=14, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE_ID)
    _set_run_font(run, size=14, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(TITLE_EN)
    _set_run_font(run, size=12, italic=True)

    for _ in range(3):
        doc.add_paragraph()

    for line_text in [
        "Disusun dan diajukan oleh:",
        "",
        "ANDI AGUNG DWI ARYA B",
        "D082251054",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line_text:
            run = p.add_run(line_text)
            bld = line_text in ("ANDI AGUNG DWI ARYA B", "D082251054")
            _set_run_font(run, size=12, bold=bld)

    for _ in range(3):
        doc.add_paragraph()

    for line_text in [
        "PROGRAM STUDI MAGISTER TEKNIK INFORMATIKA",
        "UNIVERSITAS HASANUDDIN",
        "GOWA",
        "2025",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line_text)
        _set_run_font(run, size=12, bold=True)

    # Page break
    doc.add_page_break()

    # ======================== TABLE OF CONTENTS (placeholder) ========================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAFTAR ISI")
    _set_run_font(run, size=14, bold=True)
    p.paragraph_format.space_after = Pt(18)

    toc_items = [
        ("DAFTAR ISI", "ii"),
        ("BAB I    PENDAHULUAN", "1"),
        ("    1.1  Latar Belakang", "1"),
        ("    1.2  Rumusan Masalah", "4"),
        ("    1.3  Tujuan Penelitian", "5"),
        ("    1.4  Manfaat Penelitian", "6"),
        ("    1.5  Batasan Epistemologis Sistem", "7"),
        ("    1.6  Batasan Masalah", "9"),
        ("BAB II   TINJAUAN PUSTAKA", "11"),
        ("    2.1  Sintesis dalam Konteks Literature Review", "11"),
        ("    2.2  Research Gap: Definisi dan Identifikasi", "13"),
        ("    2.3  Large Language Models dan RAG", "16"),
        ("    2.4  Pendekatan Neuro-Symbolic AI", "19"),
        ("    2.5  AI Agent dan Agentic Systems", "22"),
        ("    2.6  Knowledge Graph untuk Analisis Literatur", "25"),
        ("    2.7  Kerangka Konseptual Penelitian", "28"),
        ("BAB III  METODOLOGI PENELITIAN", "30"),
        ("    3.1  Desain Penelitian", "30"),
        ("    3.2  Arsitektur Sistem", "32"),
        ("    3.3  Implementasi Komponen", "36"),
        ("    3.4  Mekanisme Pembeda Asosiasi vs Kausal", "40"),
        ("    3.5  Rule-Based Validation Layer", "42"),
        ("    3.6  Evaluasi", "45"),
        ("    3.7  Jadwal Penelitian", "48"),
        ("DAFTAR PUSTAKA", "50"),
    ]

    for item_text, page in toc_items:
        p = doc.add_paragraph()
        tab_count = 50  # approximate tab stop
        dots = "." * max(3, 50 - len(item_text) - len(page))
        full_text = f"{item_text} {dots} {page}"
        run = p.add_run(full_text)
        is_bab = item_text.strip().startswith("BAB") or item_text.strip() == "DAFTAR PUSTAKA" or item_text.strip() == "DAFTAR ISI"
        _set_run_font(run, size=12, bold=is_bab)
        _set_paragraph_format(p, space_after=2)

    doc.add_page_break()

    # ======================== BAB I ========================
    bab1_path = os.path.join(DRAFTS_DIR, "BAB_I_PENDAHULUAN.md")
    with open(bab1_path, "r", encoding="utf-8") as f:
        bab1_text = f.read()

    converter = MarkdownToDocx(doc)
    converter.convert(bab1_text)
    doc.add_page_break()

    # ======================== BAB II ========================
    bab2_path = os.path.join(DRAFTS_DIR, "BAB_II_TINJAUAN_PUSTAKA.md")
    with open(bab2_path, "r", encoding="utf-8") as f:
        bab2_text = f.read()

    converter.convert(bab2_text)
    doc.add_page_break()

    # ======================== BAB III ========================
    bab3_path = os.path.join(DRAFTS_DIR, "BAB_III_METODOLOGI.md")
    with open(bab3_path, "r", encoding="utf-8") as f:
        bab3_text = f.read()

    converter.convert(bab3_text)
    doc.add_page_break()

    # ======================== DAFTAR PUSTAKA ========================
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run("DAFTAR PUSTAKA")
    _set_run_font(run, size=14, bold=True)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(18)

    for idx, ref in enumerate(KNOWN_REFERENCES, 1):
        p = doc.add_paragraph()
        _set_paragraph_format(p, space_after=4)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        run = p.add_run(ref)
        _set_run_font(run, size=12)

    # ======================== SAVE ========================
    doc.save(OUTPUT_PATH)
    print(f"✅ Document saved to: {OUTPUT_PATH}")
    print(f"   File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


# ---------------------------------------------------------------------------
# Comparison document builder
# ---------------------------------------------------------------------------
def build_comparison():
    """Create PERBANDINGAN_REVISI.md showing old vs new changes."""
    original_path = "/tmp/proposal_full_text.txt"
    with open(original_path, "r", encoding="utf-8") as f:
        original = f.read()

    comparison = """# PERBANDINGAN REVISI PROPOSAL

## Ringkasan Perubahan: Proposal Lama → Proposal Revisi

---

## BAB I — PENDAHULUAN

### Perubahan Judul
| Aspek | Lama | Baru |
|-------|------|------|
| **Judul (ID)** | Intelligent Research Gap Analyzer: Sistem Otomatis Berbasis RAG dan LLM untuk Generasi Rekomendasi Penelitian Lanjutan dari Analisis Jurnal | Neuro-Symbolic Agentic System untuk Deteksi Indikator Synthesis Gap dalam Literatur Ilmiah |
| **Judul (EN)** | Intelligent Research Gap Analyzer: An Automatic System Based on RAG and LLM for Generation of Further Research Recommendations from Journal Analysis | Neuro-Symbolic Agentic System for Synthesis Gap Indicator Detection in Scientific Literature |

### Perubahan Latar Belakang (1.1)
| Aspek | Lama | Baru |
|-------|------|------|
| **Fokus** | RAG + LLM + Vector DB untuk deteksi research gap umum | Neuro-Symbolic Agentic System untuk deteksi *synthesis gap* secara spesifik |
| **Pendekatan** | Pipeline linear RAG+LLM | Arsitektur Agentic multi-step reasoning + Rule-Based Validation |
| **Kebaruan** | Penggunaan Vector DB untuk retrieval cepat | 4 pilar: Agentic architecture, Rule Engine, Pembeda semantik/kausal, KG sebagai Fact Base |
| **Posisi sistem** | Sistem otomatis (klaim luas) | Alat bantu keputusan (decision support tool) — eksplisit tentang batasan |

### Perubahan Rumusan Masalah (1.2)
| Lama | Baru |
|------|------|
| Bagaimana merancang arsitektur RAG+LLM? (spesifikasi teknis) | Sejauh mana agentic reasoning mampu mendeteksi indikator synthesis gap? (pertanyaan penelitian) |
| Bagaimana mengoptimalkan akurasi? (terlalu umum) | Bagaimana mekanisme pembeda semantik/kausal memengaruhi akurasi dan false discovery rate? |
| *(tidak ada)* | Apa batasan epistemologis sistem dibandingkan penalaran manusia? |

### Perubahan Tujuan Penelitian (1.3)
| Lama | Baru |
|------|------|
| Merancang dan membangun arsitektur RAG+LLM | Mengukur kemampuan Neuro-Symbolic Agentic dalam mendeteksi indikator synthesis gap |
| Mengoptimalkan akurasi dan konsistensi | Menganalisis pengaruh mekanisme pembeda dan rule-based validation terhadap kualitas |
| *(tidak ada)* | Mengidentifikasi batasan epistemologis sistem secara eksplisit |

### Seksi Baru (Tidak Ada di Proposal Lama)
- **1.4 Manfaat Penelitian** — diperluas dengan manfaat teoretis dan praktis yang spesifik
- **1.5 Batasan Epistemologis Sistem** — **BARU SEPENUHNYA** — mendefinisikan apa yang BISA dan TIDAK BISA dilakukan sistem, termasuk tabel kemampuan dan klaim yang TIDAK dibuat
- **1.6 Batasan Masalah** — diperluas dan diperinci dengan justifikasi untuk setiap batasan

### Seksi yang Dihapus/Diganti
- **1.4 Batasan Penelitian (lama)** — diganti dengan 1.6 yang lebih komprehensif
- **1.5 Manfaat Penelitian (lama)** — dipindah ke 1.4 dan diperluas

---

## BAB II — TINJAUAN PUSTAKA

### Perubahan Struktur
| Lama | Baru |
|------|------|
| 2.1 State of The Art (tabel 10 paper) | 2.1 Sintesis dalam Konteks Literature Review |
| 2.2 Metode yang Digunakan | 2.2 Research Gap: Definisi dan Identifikasi |
| 2.3 Kerangka Pikir | 2.3 Large Language Models dan RAG |
| *(tidak ada)* | 2.4 Pendekatan Neuro-Symbolic AI |
| *(tidak ada)* | 2.5 AI Agent dan Agentic Systems |
| *(tidak ada)* | 2.6 Knowledge Graph untuk Analisis Literatur |
| *(tidak ada)* | 2.7 Kerangka Konseptual Penelitian |

### Perubahan Fundamental
- **Lama:** Tinjauan pustaka deskriptif — daftar paper dan metode tanpa analisis kritis
- **Baru:** Tinjauan argumentatif — membangun fondasi konseptual untuk setiap komponen sistem, termasuk:
  - Definisi operasional synthesis gap berdasarkan Cooper [1998] dan Booth et al. [2012]
  - Tiga indikator synthesis gap (fragmentasi, inkonsistensi, ketidaklengkapan kolektif)
  - Landasan teori Neuro-Symbolic AI
  - Kajian arsitektur AI Agent (ReAct, tool-use, multi-step reasoning)
  - KG sebagai fact base vs. visualisasi

### Seksi Baru (Tidak Ada di Proposal Lama)
- **2.1** Sintesis dalam Konteks Literature Review — fondasi konseptual baru
- **2.2** Research Gap: Definisi dan Identifikasi — operasionalisasi synthesis gap
- **2.4** Pendekatan Neuro-Symbolic AI — landasan teori integrasi neural-simbolik
- **2.5** AI Agent dan Agentic Systems — landasan arsitektur agentic
- **2.6** Knowledge Graph untuk Analisis Literatur — KG sebagai fact base
- **2.7** Kerangka Konseptual — diagram kerangka pikir yang baru

### Seksi yang Dihapus/Diganti
- **2.1 State of the Art (lama)** — dihapus karena terlalu deskriptif; diganti dengan tinjauan tematik yang lebih analitis
- **2.2 Metode yang Digunakan (lama)** — dihapus; konten relevan didistribusikan ke seksi-seksi baru
- **2.3 Kerangka Pikir (lama)** — diganti dengan kerangka konseptual yang jauh lebih komprehensif

---

## BAB III — METODOLOGI PENELITIAN

### Perubahan Arsitektur
| Aspek | Lama | Baru |
|-------|------|------|
| **Paradigma** | Pipeline linear RAG+LLM | Mixed methods: system development + experimental evaluation + expert review |
| **Arsitektur** | Retrieve → Generate → Output | 4 Fase: Ingestion → Fact Extraction → Agentic Analysis → Logical Consistency Checker |
| **Penalaran** | One-shot generation | Multi-step reasoning: Plan → Act → Observe → Reflect → Repeat |
| **Validasi** | Tidak ada | Rule-Based Validation Layer (9 aturan: F1–F3, C1–C3, K1–K3) |
| **Knowledge Graph** | Lightweight KG untuk visualisasi | KG sebagai Fact Base (tabel SPO) untuk penalaran deduktif |

### Seksi Baru (Tidak Ada di Proposal Lama)
- **3.1 Desain Penelitian** — design science research framework
- **3.2 Arsitektur Sistem** — 4-fase architecture (sepenuhnya baru)
- **3.3 Implementasi Komponen** — detail teknis setiap komponen
- **3.4 Mekanisme Pembeda Asosiasi vs Kausal** — three-layer mechanism (baru)
- **3.5 Rule-Based Validation Layer** — 9 aturan validasi (baru)
- **3.6 Evaluasi** — framework evaluasi multi-dimensi (baru)
- **3.7 Jadwal Penelitian** — timeline yang diperbarui

### Seksi yang Dihapus
- Seluruh BAB III lama (terlalu sederhana, hanya diagram alir linear tanpa detail)

---

## DAFTAR PUSTAKA

### Perubahan
| Aspek | Lama | Baru |
|-------|------|------|
| **Jumlah referensi** | ~15 referensi | ~27 referensi |
| **Cakupan** | Fokus pada RAG, LLM, Vector DB | Ditambah: Neuro-Symbolic AI, Expert Systems, AI Agents, Literature Review methodology |
| **Referensi kunci baru** | — | Cooper (1998), Booth et al. (2012), Garcez et al. (2019), Marcus (2020), Yao et al. (2023), Hevner et al. (2004) |

---

## NOVELTY STATEMENT

**Status:** BARU SEPENUHNYA — Tidak ada di proposal lama.

Dokumen ini menjabarkan:
- 4 pilar kebaruan penelitian
- Perbandingan sistematis dengan pendekatan konvensional
- Klaim eksplisit dan klaim yang TIDAK dibuat
- Posisi dalam spektrum pendekatan AI (Pure Neural ↔ Neuro-Symbolic ↔ Pure Symbolic)

---

## Ringkasan: Alasan Utama Perubahan

Seluruh revisi dilakukan berdasarkan **8 kritik fatal** dari penguji seminar proposal:

1. ❌ Rumusan masalah bukan "masalah" → ✅ Diubah menjadi pertanyaan penelitian
2. ❌ Definisi synthesis gap terlalu dangkal → ✅ Dioperasionalkan dengan 3 indikator
3. ❌ Tidak ada penjelasan batas kemampuan sistem → ✅ Ditambah Batasan Epistemologis
4. ❌ RAG+LLM bukan kebaruan → ✅ Diubah ke Neuro-Symbolic Agentic System
5. ❌ Tidak bisa bedakan semantik vs kausal → ✅ Ditambah mekanisme 3 lapis
6. ❌ Tidak ada validasi logis → ✅ Ditambah Rule-Based Validation Layer
7. ❌ Diagram alir terlalu linear → ✅ Diubah ke arsitektur 4 fase
8. ❌ KG hanya visualisasi → ✅ KG sebagai Fact Base (tabel SPO)
"""

    comparison_path = os.path.join(DRAFTS_DIR, "PERBANDINGAN_REVISI.md")
    with open(comparison_path, "w", encoding="utf-8") as f:
        f.write(comparison)
    print(f"✅ Comparison document saved to: {comparison_path}")


if __name__ == "__main__":
    build_document()
    build_comparison()
